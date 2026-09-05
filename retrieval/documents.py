from __future__ import annotations

import io
import re
import unicodedata
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader


@dataclass
class ParsedDocument:
    """Normalized document returned by the ingestion pipeline."""

    doc_id: str
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)
    sections: dict[str, str] = field(default_factory=dict)
    error: str | None = None


SECTION_NAMES = (
    "Candidate Skills",
    "Work Experience",
    "Education",
    "JD Requirements",
    "Core Competencies",
)

_SECTION_ALIASES = {
    "candidate skills": "Candidate Skills",
    "skills": "Candidate Skills",
    "technical skills": "Candidate Skills",
    "work experience": "Work Experience",
    "experience": "Work Experience",
    "professional experience": "Work Experience",
    "education": "Education",
    "academic background": "Education",
    "jd requirements": "JD Requirements",
    "job requirements": "JD Requirements",
    "requirements": "JD Requirements",
    "job description": "JD Requirements",
    "core competencies": "Core Competencies",
    "competencies": "Core Competencies",
}

_NON_PRINTABLE_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_WHITESPACE_RE = re.compile(r"[ \t]+")
_BLANK_LINES_RE = re.compile(r"\n{3,}")


def _new_document(
    text: str,
    source_type: str,
    file_name: str | None = None,
) -> ParsedDocument:
    """Normalize text and build the standardized document result."""
    normalized = _sanitize_text(text)
    sections = _parse_sections(normalized)

    metadata: dict[str, Any] = {
        "source_type": source_type,
        "encoding": "utf-8",
    }

    if file_name:
        metadata["file_name"] = file_name

    return ParsedDocument(
        doc_id=uuid.uuid4().hex,
        text=normalized,
        metadata=metadata,
        sections=sections,
    )


def _error_document(
    error: str,
    source_type: str = "unknown",
    file_name: str | None = None,
) -> ParsedDocument:
    """Return a structured error result instead of raising ingestion errors."""
    metadata: dict[str, Any] = {
        "source_type": source_type,
        "encoding": "utf-8",
    }

    if file_name:
        metadata["file_name"] = file_name

    return ParsedDocument(
        doc_id=uuid.uuid4().hex,
        text="",
        metadata=metadata,
        sections={},
        error=error,
    )


def _sanitize_text(text: str) -> str:
    """Remove encoding artifacts, control characters, and redundant whitespace."""
    text = unicodedata.normalize("NFKC", text)

    replacements = {
        "\ufeff": "",
        "\ufffd": "",
        "\u00a0": " ",
        "\u200b": "",
        "\u200c": "",
        "\u200d": "",
        "\u2060": "",
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    text = _NON_PRINTABLE_RE.sub("", text)

    lines = []
    for line in text.splitlines():
        line = _WHITESPACE_RE.sub(" ", line).strip()
        if line:
            lines.append(line)

    return _BLANK_LINES_RE.sub("\n\n", "\n".join(lines)).strip()


def _normalize_heading(value: str) -> str:
    """Normalize a possible section heading for alias matching."""
    value = value.strip().rstrip(":")
    value = re.sub(r"[^a-zA-Z0-9 ]+", " ", value)
    value = re.sub(r"\s+", " ", value)
    return value.lower().strip()


def _parse_sections(text: str) -> dict[str, str]:
    """Extract known resume/JD sections using heading-based parsing."""
    if not text:
        return {}

    lines = text.splitlines()
    sections: dict[str, list[str]] = {}
    current_section: str | None = None

    for line in lines:
        heading = _SECTION_ALIASES.get(_normalize_heading(line))

        if heading:
            current_section = heading
            sections.setdefault(heading, [])
            continue

        if current_section:
            sections[current_section].append(line)

    return {
        name: "\n".join(values).strip()
        for name, values in sections.items()
        if "\n".join(values).strip()
    }


def _read_text_file(path: Path) -> tuple[str, str]:
    """Read a text file with a UTF-8-first decoding strategy."""
    raw = path.read_bytes()

    if not raw:
        raise ValueError("The input file is empty.")

    try:
        return raw.decode("utf-8"), "text"
    except UnicodeDecodeError:
        try:
            return raw.decode("utf-8-sig"), "text"
        except UnicodeDecodeError as exc:
            raise ValueError("The input text is not valid UTF-8.") from exc


def _read_pdf(source: bytes) -> str:
    """Extract text from a PDF byte stream."""
    if not source:
        raise ValueError("The input file is empty.")

    try:
        reader = PdfReader(io.BytesIO(source))
    except Exception as exc:
        raise ValueError("The PDF file is malformed or cannot be read.") from exc

    if reader.is_encrypted:
        try:
            result = reader.decrypt("")
        except Exception as exc:
            raise ValueError("The PDF is password-protected.") from exc

        if result == 0:
            raise ValueError("The PDF is password-protected.")

    pages = []

    try:
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text:
                pages.append(page_text)
    except Exception as exc:
        raise ValueError("The PDF text could not be extracted.") from exc

    text = "\n".join(pages)

    if not text.strip():
        raise ValueError("The PDF contains no extractable text.")

    return text


def _read_docx(source: bytes) -> str:
    """Extract paragraph and table text from a DOCX byte stream."""
    if not source:
        raise ValueError("The input file is empty.")

    try:
        document = Document(io.BytesIO(source))
    except Exception as exc:
        raise ValueError("The DOCX file is malformed or cannot be read.") from exc

    parts = [
        paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()
    ]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                parts.append(row_text)

    text = "\n".join(parts)

    if not text.strip():
        raise ValueError("The DOCX contains no extractable text.")

    return text


def _load_document(
    file_path_or_bytes: str | Path | bytes,
    document_type: str,
) -> ParsedDocument:
    """Load a document from a path or bytes and return a structured result."""
    if isinstance(file_path_or_bytes, bytes):
        source = file_path_or_bytes
        file_name = None
        source_type = document_type

        if not source:
            return _error_document(
                "The input file is empty.",
                source_type=document_type,
            )

        try:
            if source.startswith(b"%PDF"):
                text = _read_pdf(source)
                source_type = "pdf"
            elif source.startswith(b"PK"):
                text = _read_docx(source)
                source_type = "docx"
            else:
                try:
                    text = source.decode("utf-8")
                except UnicodeDecodeError as exc:
                    raise ValueError("The input text is not valid UTF-8.") from exc
                source_type = document_type

            return _new_document(text, source_type, file_name)

        except (ValueError, OSError) as exc:
            return _error_document(str(exc), source_type, file_name)

    path = Path(file_path_or_bytes)
    file_name = path.name

    if not path.exists():
        return _error_document(
            f"File not found: {path}",
            document_type,
            file_name,
        )

    if not path.is_file():
        return _error_document(
            f"Input path is not a file: {path}",
            document_type,
            file_name,
        )

    suffix = path.suffix.lower()

    try:
        if suffix == ".pdf":
            text = _read_pdf(path.read_bytes())
            source_type = "pdf"
        elif suffix == ".docx":
            text = _read_docx(path.read_bytes())
            source_type = "docx"
        else:
            text, source_type = _read_text_file(path)

        return _new_document(text, source_type, file_name)

    except (ValueError, OSError) as exc:
        return _error_document(str(exc), document_type, file_name)


def load_resume(file_path_or_bytes: str | Path | bytes) -> ParsedDocument:
    """Load and normalize a candidate resume."""
    return _load_document(file_path_or_bytes, "resume")


def load_jd(file_path_or_bytes: str | Path | bytes) -> ParsedDocument:
    """Load and normalize a job description."""
    return _load_document(file_path_or_bytes, "job_description")
