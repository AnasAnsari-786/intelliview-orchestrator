import io

from docx import Document
from reportlab.pdfgen import canvas

from retrieval.documents import ParsedDocument, load_jd, load_resume


def test_load_resume_from_utf8_bytes():
    data = b"""
    Candidate Skills:
    Python, FastAPI, SQL

    Work Experience:
    Software Developer at Example Corp

    Education:
    B.Tech Computer Engineering
    """

    result = load_resume(data)

    assert isinstance(result, ParsedDocument)
    assert result.error is None
    assert result.text
    assert result.metadata["source_type"] == "resume"
    assert result.sections["Candidate Skills"] == "Python, FastAPI, SQL"
    assert "Software Developer" in result.sections["Work Experience"]
    assert "B.Tech Computer Engineering" in result.sections["Education"]


def test_load_jd_from_utf8_bytes():
    data = b"""
    JD Requirements:
    Python, FastAPI, REST APIs

    Core Competencies:
    Problem solving, communication, teamwork
    """

    result = load_jd(data)

    assert isinstance(result, ParsedDocument)
    assert result.error is None
    assert result.metadata["source_type"] == "job_description"
    assert "Python, FastAPI, REST APIs" in result.sections["JD Requirements"]
    assert "Problem solving" in result.sections["Core Competencies"]


def test_whitespace_and_unicode_artifacts_are_sanitized():
    data = (
        "Candidate Skills:\n"
        "Python\u00a0  FastAPI   SQL\n"
        "\u200bDeveloper\u200d Experience"
    ).encode()

    result = load_resume(data)

    assert result.error is None
    assert "\u00a0" not in result.text
    assert "\u200b" not in result.text
    assert "\u200d" not in result.text
    assert "Python FastAPI SQL" in result.text


def test_docx_resume_loading():
    document = Document()
    document.add_paragraph("Candidate Skills:")
    document.add_paragraph("Python, Machine Learning")
    document.add_paragraph("Work Experience:")
    document.add_paragraph("AI Engineer - 2 years")
    document.add_paragraph("Education:")
    document.add_paragraph("B.Tech Computer Engineering")

    buffer = io.BytesIO()
    document.save(buffer)

    result = load_resume(buffer.getvalue())

    assert result.error is None
    assert result.metadata["source_type"] == "docx"
    assert "Python, Machine Learning" in result.sections["Candidate Skills"]
    assert "AI Engineer - 2 years" in result.sections["Work Experience"]
    assert "B.Tech Computer Engineering" in result.sections["Education"]


def test_pdf_resume_loading():
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.drawString(72, 750, "Candidate Skills:")
    pdf.drawString(72, 730, "Python, FastAPI")
    pdf.drawString(72, 700, "Work Experience:")
    pdf.drawString(72, 680, "Backend Developer")
    pdf.drawString(72, 650, "Education:")
    pdf.drawString(72, 630, "B.Tech Computer Engineering")
    pdf.save()

    result = load_resume(buffer.getvalue())

    assert result.error is None
    assert result.metadata["source_type"] == "pdf"
    assert "Python, FastAPI" in result.sections["Candidate Skills"]
    assert "Backend Developer" in result.sections["Work Experience"]
    assert "B.Tech Computer Engineering" in result.sections["Education"]


def test_empty_input_returns_structured_error():
    result = load_resume(b"")

    assert isinstance(result, ParsedDocument)
    assert result.error is not None
    assert result.text == ""
    assert result.sections == {}


def test_malformed_utf8_returns_structured_error():
    result = load_resume(b"\xff\xfe\xfa")

    assert isinstance(result, ParsedDocument)
    assert result.error is not None
    assert "UTF-8" in result.error


def test_malformed_pdf_returns_structured_error():
    result = load_resume(b"%PDF-this-is-not-a-valid-pdf")

    assert isinstance(result, ParsedDocument)
    assert result.error is not None


def test_missing_file_returns_structured_error():
    result = load_resume("does-not-exist.pdf")

    assert isinstance(result, ParsedDocument)
    assert result.error is not None
    assert result.text == ""


def test_doc_ids_are_unique():
    first = load_resume(b"Candidate Skills:\nPython")
    second = load_resume(b"Candidate Skills:\nPython")

    assert first.doc_id
    assert second.doc_id
    assert first.doc_id != second.doc_id
